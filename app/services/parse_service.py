"""文档文本提取与解析结果落库。"""

from dataclasses import dataclass
from pathlib import Path

import fitz
import numpy as np
from docx import Document
from PIL import Image
from sqlalchemy import select

from app.config import get_settings
from app.db import get_session
from app.models.parse import ContractParse
from app.services.ocr_service import filter_low_confidence, ocr_image


@dataclass(slots=True)
class DocumentPage:
    """文档页。"""

    page_no: int
    text: str


@dataclass(slots=True)
class DocumentText:
    """文档文本。"""

    file_path: str
    pages: list[DocumentPage]


def _text_from_docx(file_path: str) -> str:
    """提取 docx 段落与表格文本。"""
    doc = Document(file_path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_docx(file_path: str) -> DocumentText:
    """Word 文本提取（段落 + 表格），docx 无真实分页记 page_no=1。"""
    text = _text_from_docx(file_path)
    if not text.strip():
        raise ValueError("文档内容为空")
    return DocumentText(file_path=str(file_path), pages=[DocumentPage(page_no=1, text=text)])


def parse_pdf(file_path: str) -> DocumentText:
    """PDF 逐页判断文本层：有层走 pdfplumber 逻辑（PyMuPDF 取文本），无层转图片 OCR。"""
    settings = get_settings()
    doc = fitz.open(file_path)
    if doc.page_count > settings.max_pdf_pages:
        raise ValueError(f"PDF 页数超过上限 {settings.max_pdf_pages}")
    pages: list[DocumentPage] = []
    for index in range(doc.page_count):
        page = doc[index]
        text = page.get_text().strip()
        if text:
            pages.append(DocumentPage(page_no=index + 1, text=text))
            continue
        pix = page.get_pixmap(dpi=150)
        image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        lines = filter_low_confidence(ocr_image(np.array(image)), settings.ocr_confidence_threshold)
        if not lines:
            raise ValueError(f"PDF 第 {index + 1} 页图片无法识别")
        pages.append(DocumentPage(page_no=index + 1, text="\n".join(line.text for line in lines)))
    if not pages:
        raise ValueError("文档内容为空")
    return DocumentText(file_path=str(file_path), pages=pages)


def parse_image(file_path: str) -> DocumentText:
    """图片 OCR 提取文本。"""
    settings = get_settings()
    image = Image.open(file_path)
    lines = filter_low_confidence(ocr_image(np.array(image)), settings.ocr_confidence_threshold)
    if not lines:
        raise ValueError("图片无法识别")
    return DocumentText(
        file_path=str(file_path),
        pages=[DocumentPage(page_no=1, text="\n".join(line.text for line in lines))],
    )


def extract_document_text(file_path: str) -> DocumentText:
    """按扩展名分发文档解析。"""
    suffix = Path(file_path).suffix.lower()
    if suffix == ".docx":
        return parse_docx(file_path)
    if suffix == ".pdf":
        return parse_pdf(file_path)
    if suffix in {".png", ".jpg", ".jpeg"}:
        return parse_image(file_path)
    raise ValueError(f"不支持的附件类型: {suffix}")


def merge_documents(texts: list[DocumentText]) -> DocumentText:
    """按附件顺序合并全部附件，页号连续编号。"""
    pages: list[DocumentPage] = []
    page_no = 1
    for document_text in texts:
        for page in document_text.pages:
            pages.append(DocumentPage(page_no=page_no, text=page.text))
            page_no += 1
    return DocumentText(file_path="merged", pages=pages)


def save_parse_result(task_id: int, parse_result: dict):
    """写 contract_parses，更新 parse_status。"""
    with next(get_session()) as db:
        record = db.scalar(select(ContractParse).where(ContractParse.task_id == task_id))
        if record is None:
            record = ContractParse(task_id=task_id)
            db.add(record)
        record.basic_info_json = parse_result.get("basic_info") or {}
        record.clause_info_json = parse_result.get("clauses") or {}
        record.parse_status = parse_result.get("parse_status", "success")
        record.parse_error = parse_result.get("parse_error")
        db.commit()


def mark_parse_failed(task_id: int, parse_error: str):
    """记录 parse_error 并置 parse_status=failed。"""
    with next(get_session()) as db:
        record = db.scalar(select(ContractParse).where(ContractParse.task_id == task_id))
        if record is None:
            record = ContractParse(task_id=task_id)
            db.add(record)
        record.parse_status = "failed"
        record.parse_error = parse_error
        db.commit()


def get_parse_by_task(task_id: int) -> ContractParse | None:
    """查询解析结果。"""
    with next(get_session()) as db:
        return db.scalar(select(ContractParse).where(ContractParse.task_id == task_id))
