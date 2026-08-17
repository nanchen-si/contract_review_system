"""生成 mock 合同附件：三份 docx 与一份图片扫描件。"""

from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parent.parent
ATTACH_DIR = ROOT / "mock" / "attachments"


def _write_docx(path: Path, title: str, clauses: list[str]):
    """按给定条款生成 docx 合同。"""
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph("合同编号：HT-2026-0088")
    doc.add_paragraph("甲方（签约主体）：北京星辰科技有限公司")
    doc.add_paragraph("乙方（对方名称）：深圳启明电子有限公司")
    for clause in clauses:
        doc.add_paragraph(clause)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def gen_high_risk_contract() -> Path:
    """生成高风险合同 docx：预付款 80%、自动续约、无保密条款。"""
    path = ATTACH_DIR / "高风险采购合同.docx"
    _write_docx(
        path,
        "高风险采购合同",
        [
            "第一条 合同金额：合同总金额为人民币 1,200,000 元，币种为人民币。",
            "第二条 付款条款：合同签订后 10 个工作日内支付预付款 80%。",
            "第三条 合同期限：本合同自 2026-08-20 起生效，有效期一年，到期自动续约。",
            "第四条 违约责任：任何一方违约，应按合同金额的 10% 支付违约金。",
        ],
    )
    return path


def gen_medium_risk_contract() -> Path:
    """生成中风险合同 docx：缺验收标准、境外管辖地。"""
    path = ATTACH_DIR / "中风险采购合同.docx"
    _write_docx(
        path,
        "中风险采购合同",
        [
            "第一条 合同金额：合同总金额为人民币 800,000 元，币种为人民币。",
            "第二条 付款条款：合同签订后支付预付款 30%，验收合格后 30 日内付清。",
            "第三条 交付条款：乙方应于 2026-12-31 前完成全部货物交付。",
            "第四条 争议解决条款：双方争议由香港法院管辖。",
        ],
    )
    return path


def gen_low_risk_contract() -> Path:
    """生成低风险合同 docx：条款齐全，作为 demo 主合同。"""
    path = ATTACH_DIR / "采购合同2026.docx"
    _write_docx(
        path,
        "供应商采购合同",
        [
            "第一条 合同金额：合同总金额为人民币 1,200,000 元，币种为人民币。",
            "第二条 付款条款：合同签订后 10 个工作日内支付预付款 60%；验收合格后 30 日内支付剩余 40%。",
            "第三条 交付条款：乙方应于 2026-12-31 前完成全部货物交付。",
            "第四条 验收条款：甲方应在乙方交付后 15 个工作日内组织验收，验收标准以双方确认的技术协议为准。",
            "第五条 违约条款：任何一方违约，应按合同金额的 10% 支付违约金。",
            "第六条 保密条款：双方应对合同内容及商业秘密承担保密义务，保密期限为合同终止后三年。",
            "第七条 数据条款：乙方处理甲方数据应遵守数据安全法及相关规定。",
            "第八条 知识产权条款：软件相关知识产权归甲方所有。",
            "第九条 争议解决条款：双方争议由合同签订地人民法院管辖。",
        ],
    )
    return path


def gen_scanned_contract() -> Path:
    """把中风险合同内容渲染为 PNG 扫描件，供 OCR 演示。"""
    path = ATTACH_DIR / "采购合同2026-扫描件.png"
    path.parent.mkdir(parents=True, exist_ok=True)
    lines = [
        "中风险采购合同（扫描件）",
        "合同编号：HT-2026-0088",
        "甲方：北京星辰科技有限公司  乙方：深圳启明电子有限公司",
        "第一条 合同金额：人民币 800,000 元，币种为人民币。",
        "第二条 付款条款：预付款 30%，验收合格后 30 日内付清。",
        "第三条 交付条款：乙方应于 2026-12-31 前完成全部货物交付。",
        "第四条 争议解决条款：双方争议由香港法院管辖。",
        "（本合同未包含验收标准条款）",
    ]
    image = Image.new("RGB", (1200, 900), "white")
    draw = ImageDraw.Draw(image)
    font = None
    for font_path in ("C:/Windows/Fonts/msyh.ttc", "C:/Windows/Fonts/simhei.ttf"):
        if Path(font_path).exists():
            font = ImageFont.truetype(font_path, 30)
            break
    y = 60
    for line in lines:
        draw.text((60, y), line, fill="black", font=font)
        y += 70
    image.save(path)
    return path


def main():
    """依次生成并校验文件存在。"""
    files = [
        gen_high_risk_contract(),
        gen_medium_risk_contract(),
        gen_low_risk_contract(),
        gen_scanned_contract(),
    ]
    for file in files:
        assert file.exists(), f"missing: {file}"
        print(file)


if __name__ == "__main__":
    main()
