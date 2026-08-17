"""文档解析服务单元测试。"""

from docx import Document

from app.services.parse_service import (
    extract_document_text,
    merge_documents,
    parse_docx,
)


def _make_docx(tmp_path):
    path = tmp_path / "合同.docx"
    doc = Document()
    doc.add_paragraph("合同编号：HT-2026-0088")
    doc.add_paragraph("第三条 付款条款：预付款 60%")
    doc.save(path)
    return str(path)


def test_docx_extract(tmp_path):
    """docx 文本提取。"""
    text = parse_docx(_make_docx(tmp_path))
    assert text.pages[0].page_no == 1
    assert "合同编号" in text.pages[0].text


def test_extract_document_text(tmp_path):
    """按扩展名分发。"""
    text = extract_document_text(_make_docx(tmp_path))
    assert text.pages


def test_merge_documents(tmp_path):
    """多附件合并连续编页。"""
    one = parse_docx(_make_docx(tmp_path))
    merged = merge_documents([one, one])
    assert [page.page_no for page in merged.pages] == [1, 2]
